import os
import re
import base64
import PyPDF2
from collections import Counter
from docx import Document
import requests
from groq import Groq
import json

GROQ_API_KEY = "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"
GITHUB_TOKEN = "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"

groq_client = Groq(api_key=GROQ_API_KEY)

CODE_EXTENSIONS = {
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".cpp", ".c", ".cs",
    ".go", ".rs", ".swift", ".kt", ".sql", ".ipynb", ".rb", ".sh", ".html", ".css"
}
SKIP_PATTERNS = [
    "package-lock.json", "yarn.lock", ".gitignore", "node_modules",
    "dist/", "build/", "__pycache__", ".min.js", "vendor/", "migrations/"
]


class Agent1:

    MODELS = [
        "llama-3.3-70b-versatile",
        "llama3-70b-8192",
        "mixtral-8x7b-32768",
        "gemma2-9b-it",
    ]

    def __init__(self, file_path):
        self.file_path = file_path
        self.headers = {
            "Authorization": f"token {GITHUB_TOKEN}",
            "Accept": "application/vnd.github.v3+json"
        }

    # ── Resume reading ────────────────────────────────────────────────────────

    def read_file(self):
        ext = os.path.splitext(self.file_path)[1].lower()
        if ext == ".pdf":
            with open(self.file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                return "\n".join([p.extract_text() or "" for p in reader.pages])
        elif ext == ".docx":
            doc = Document(self.file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            return "Please upload a PDF or Word file."

    def extract_links_from_word(self):
        links = []
        doc = Document(self.file_path)
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype:
                links.append(rel.target_ref)
        return list(set(links))

    def extract_links_from_pdf(self):
        links = []
        with open(self.file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                if "/Annots" in page:
                    for annot in page["/Annots"]:
                        annot_obj = annot.get_object()
                        if "/A" in annot_obj and "/URI" in annot_obj["/A"]:
                            links.append(annot_obj["/A"]["/URI"])
        return list(set(links))

    def extract_links_from_text(self):
        text = self.read_file()
        return list(set(re.findall(r'https?://\S+', text)))

    def get_hyperlinks(self):
        ext = os.path.splitext(self.file_path)[1].lower()
        if ext == ".pdf":
            return list(set(self.extract_links_from_pdf() + self.extract_links_from_text()))
        elif ext == ".docx":
            return list(set(self.extract_links_from_word() + self.extract_links_from_text()))
        else:
            return []

    def display_links(self):
        links = self.get_hyperlinks()
        print("\n===== HYPERLINKS =====\n")
        if links:
            for link in links:
                print(link)
        else:
            print("No hyperlinks found.")

    def display(self):
        text = self.read_file()
        links = self.get_hyperlinks()
        combined_output = text
        if links:
            combined_output += "\n\n===== LINKS =====\n"
            combined_output += "\n".join(links)
        return combined_output

    def extract_github_link(self, text):
        pattern = r'https?://(?:www\.)?github\.com/([A-Za-z0-9-]+)/?(?=\s|$)'
        matches = re.findall(pattern, text)
        if matches:
            return f"https://github.com/{matches[0]}"
        return None

    # ── GitHub helpers ────────────────────────────────────────────────────────

    def _gh(self, url):
        try:
            r = requests.get(url, headers=self.headers, timeout=15)
            return r.json() if r.status_code == 200 else None
        except Exception:
            return None

    def _get_languages(self, username, repo_name):
        data = self._gh(f"https://api.github.com/repos/{username}/{repo_name}/languages")
        return list(data.keys()) if isinstance(data, dict) else []

    def _get_commit_count(self, username, repo_name):
        try:
            r = requests.get(
                f"https://api.github.com/repos/{username}/{repo_name}/commits?per_page=1",
                headers=self.headers, timeout=15
            )
            if r.status_code == 200:
                if "Link" in r.headers:
                    last = r.headers["Link"].split(",")[-1]
                    return int(last.split("page=")[-1].split(">")[0])
                return len(r.json()) if isinstance(r.json(), list) else 0
        except Exception:
            pass
        return 0

    def _get_readme(self, username, repo_name):
        data = self._gh(f"https://api.github.com/repos/{username}/{repo_name}/readme")
        if isinstance(data, dict) and data.get("content"):
            try:
                return base64.b64decode(
                    data["content"].replace("\n", "")
                ).decode("utf-8", errors="ignore")[:2000]
            except Exception:
                pass
        return ""

    def _get_code_files(self, username, repo_name, max_files=6):
        tree = None
        for branch in ["main", "master", "HEAD"]:
            data = self._gh(
                f"https://api.github.com/repos/{username}/{repo_name}/git/trees/{branch}?recursive=1"
            )
            if data and "tree" in data:
                tree = data["tree"]
                break

        if not tree:
            return {}

        selected = []
        for item in tree:
            if item.get("type") != "blob":
                continue
            path = item.get("path", "")
            if any(skip in path for skip in SKIP_PATTERNS):
                continue
            if os.path.splitext(path)[1].lower() not in CODE_EXTENSIONS:
                continue
            if item.get("size", 0) > 80000:
                continue
            selected.append(path)
            if len(selected) >= max_files:
                break

        code_files = {}
        for path in selected:
            data = self._gh(
                f"https://api.github.com/repos/{username}/{repo_name}/contents/{path}"
            )
            if isinstance(data, dict) and data.get("encoding") == "base64":
                try:
                    content = base64.b64decode(
                        data["content"].replace("\n", "")
                    ).decode("utf-8", errors="replace")
                    code_files[path] = content[:2000]
                except Exception:
                    pass
        return code_files

    def _get_contributor_stats(self, username, repo_name):
        data = self._gh(f"https://api.github.com/repos/{username}/{repo_name}/contributors")
        if not isinstance(data, list):
            return {}
        total = sum(c.get("contributions", 0) for c in data)
        for c in data:
            if c.get("login", "").lower() == username.lower():
                user_contribs = c.get("contributions", 0)
                pct = round((user_contribs / total * 100), 1) if total else 0
                return {
                    "user_contributions":      user_contribs,
                    "total_contributions":     total,
                    "contribution_percentage": pct
                }
        return {}

    def _get_recent_activity(self, username):
        data = self._gh(f"https://api.github.com/users/{username}/events/public?per_page=30")
        if not isinstance(data, list):
            return {}
        event_counts = Counter(e.get("type") for e in data)
        repos_active = list(set(
            e.get("repo", {}).get("name", "") for e in data if e.get("repo")
        ))
        return {
            "event_types":           dict(event_counts),
            "recently_active_repos": repos_active[:10],
            "total_recent_events":   len(data)
        }

    # ── Groq LLM helper ───────────────────────────────────────────────────────

    def _llm(self, prompt):
        """Call Groq with model fallback."""
        import time
        for model in self.MODELS:
            try:
                response = groq_client.chat.completions.create(
                    model=model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.2,
                    max_tokens=2048,
                )
                return response.choices[0].message.content.strip()
            except Exception as e:
                err = str(e)
                if "429" in err or "rate_limit" in err.lower():
                    print(f"  {model} rate limit — waiting 10s...")
                    time.sleep(10)
                    continue
                print(f"  {model} error: {e}")
                continue
        return "Could not generate summary."

    # ── Gemini-style wrappers using Groq ──────────────────────────────────────

    def _analyze_repo(self, repo_name, repo_desc, languages, stars, forks,
                       commit_count, readme_text, code_files):
        code_context = ""
        for path, content in code_files.items():
            code_context += f"\n\n--- File: {path} ---\n{content[:1500]}"

        prompt = f"""You are a senior software engineer and technical hiring expert.
Analyze this GitHub repository thoroughly for hiring assessment.

=== REPO METADATA ===
Name: {repo_name}
Description: {repo_desc}
Languages: {languages}
Stars: {stars} | Forks: {forks} | Approx Commits: {commit_count}

=== README ===
{readme_text[:1000] if readme_text else "No README"}

=== CODE FILES ===
{code_context if code_context else "No code files available"}

Return structured text with these exact sections:
1. PURPOSE — What does this project do?
2. TECH STACK — Languages, frameworks, libraries used
3. CODE QUALITY — Readability, structure, naming, modularity (score /10)
4. TECHNICAL COMPLEXITY — How complex is the code? (score /10)
5. SKILLS DEMONSTRATED — Specific skills shown in code (list each on new line with -)
6. ARCHITECTURE PATTERNS — Design patterns used
7. BEST PRACTICES — Good practices followed
8. AREAS FOR IMPROVEMENT — What is lacking
9. REAL WORLD APPLICABILITY — Production-worthy or academic?
10. RED FLAGS — Any concerns
11. OVERALL VERDICT — 1 paragraph hiring summary"""

        return self._llm(prompt)

    def _overall_candidate_assessment(self, username, profile, all_languages,
                                       all_skills, repo_summaries, activity):
        repos_text = "\n\n".join(repo_summaries[:5])

        prompt = f"""You are a senior technical recruiter and engineering manager.
Give a holistic hiring assessment based on this complete GitHub analysis.

=== CANDIDATE ===
Username: {username}
Name: {profile.get('name', '')}
Bio: {profile.get('bio', '')}
Followers: {profile.get('followers', 0)} | Public Repos: {profile.get('public_repos', 0)}

=== LANGUAGES & SKILLS ===
Languages: {', '.join(all_languages[:15])}
Skills from code: {', '.join(list(set(all_skills))[:20])}

=== RECENT ACTIVITY ===
{activity}

=== TOP REPO REVIEWS ===
{repos_text[:3000]}

Return structured assessment:
1. DEVELOPER LEVEL — Beginner / Junior / Mid-level / Senior / Expert (with reasoning)
2. TECHNICAL STRENGTHS — Top 5 with evidence
3. TECHNICAL WEAKNESSES — Top 3 with evidence
4. DOMAIN EXPERTISE — Which domains (ML, backend, frontend, DevOps, etc.)
5. CODE CONSISTENCY — Consistent quality or varies?
6. LEARNING TRAJECTORY — Is candidate growing?
7. HIRING RECOMMENDATION — Hire? For what roles? What level?
8. CONFIDENCE SCORE — 0.0 to 1.0
9. ONE LINE SUMMARY — Single sentence verdict"""

        return self._llm(prompt)

    def parse_resume(self):
        """Parse resume text into structured JSON using Groq."""
        import json
        text      = self.read_file()
        links     = self.get_hyperlinks()
        links_str = "\n".join(links) if links else "None found"

        prompt = f"""Extract all structured information from this resume.
The hyperlinks below are ACTUAL URLs embedded in the document.

Return ONLY valid JSON — no markdown, no extra text:

{{
    "name": "",
    "email": "",
    "phone": "",
    "location": "",
    "summary": "",
    "github_url": "",
    "linkedin_url": "",
    "portfolio_url": "",
    "skills": [],
    "experience": [
        {{"role": "", "company": "", "duration": "", "location": "", "description": ""}}
    ],
    "education": [
        {{"degree": "", "field": "", "institution": "", "year": ""}}
    ],
    "certifications": [],
    "projects": [
        {{"name": "", "description": "", "technologies": []}}
    ]
}}

Hyperlinks:
{links_str}

Resume text:
{text[:6000]}"""

        raw = self._llm(prompt)
        try:
            raw = raw.replace("```json", "").replace("```", "").strip()
            return json.loads(raw)
        except Exception:
            return {"raw": raw}

    # ── Main GitHub analysis ──────────────────────────────────────────────────

    def analyze_github_profile(self):
        text        = self.display()
        github_link = self.extract_github_link(text)

        if not github_link:
            return {"error": "No GitHub link found in resume."}

        username = github_link.rstrip("/").split("/")[-1]
        print(f"\n🐙 Scraping GitHub: {username}")

        profile = self._gh(f"https://api.github.com/users/{username}")
        if not profile:
            return {"error": f"GitHub user '{username}' not found."}
        print(f"  ✅ Profile: {profile.get('name')} ({profile.get('public_repos')} repos)")

        print("  📊 Fetching recent activity...")
        activity = self._get_recent_activity(username)

        repos_raw = self._gh(
            f"https://api.github.com/users/{username}/repos?per_page=100&sort=updated"
        )
        if not isinstance(repos_raw, list):
            return {"error": "Could not fetch repositories."}

        print(f"  📦 {len(repos_raw)} repos — analyzing sequentially...")

        all_languages, all_skills, repo_details, repo_summaries = [], [], [], []

        for i, repo in enumerate(repos_raw):
            repo_name = repo.get("name", "")
            repo_desc = repo.get("description", "") or ""
            stars     = repo.get("stargazers_count", 0)
            forks     = repo.get("forks_count", 0)
            is_fork   = repo.get("fork", False)
            homepage  = repo.get("homepage", "")
            updated   = repo.get("updated_at", "")
            topics    = repo.get("topics", [])

            print(f"  [{i+1}/{len(repos_raw)}] {repo_name}")

            languages    = self._get_languages(username, repo_name)
            commit_count = self._get_commit_count(username, repo_name)
            readme_text  = self._get_readme(username, repo_name)
            code_files   = self._get_code_files(username, repo_name)
            contrib      = self._get_contributor_stats(username, repo_name)

            all_languages.extend(languages)

            print(f"    🤖 LLM review...")
            llm_review = self._analyze_repo(
                repo_name, repo_desc, languages, stars, forks,
                commit_count, readme_text, code_files
            )

            skills_match = re.findall(
                r'(?:SKILLS DEMONSTRATED)[:\-\s]+(.*?)(?:\n\d+\.)',
                llm_review, re.IGNORECASE | re.DOTALL
            )
            if skills_match:
                skills = [s.strip("- \n") for s in skills_match[0].split("\n") if s.strip()]
                all_skills.extend(skills)

            repo_details.append({
                "name":                repo_name,
                "description":         repo_desc,
                "url":                 repo.get("html_url", ""),
                "languages":           languages,
                "stars":               stars,
                "forks":               forks,
                "commit_count":        commit_count,
                "last_updated":        updated,
                "is_fork":             is_fork,
                "topics":              topics,
                "has_live_demo":       bool(homepage),
                "homepage":            homepage,
                "readme_preview":      readme_text[:500],
                "code_files_analyzed": list(code_files.keys()),
                "contributor_stats":   contrib,
                "llm_review":          llm_review,
            })
            repo_summaries.append(f"Repo: {repo_name}\n{llm_review}")

        lang_counter  = Counter(all_languages)
        top_languages = [l for l, _ in lang_counter.most_common(15)]
        top_repos     = sorted(repo_details, key=lambda x: x["stars"], reverse=True)[:5]
        active_repos  = [r for r in repo_details if not r["is_fork"]]

        print("  🧠 Final candidate assessment...")
        final_assessment = self._overall_candidate_assessment(
            username, profile, top_languages,
            all_skills, repo_summaries, str(activity)
        )

        print("  ✅ GitHub analysis complete!")

        return {
            "candidate_profile": {
                "username":     username,
                "name":         profile.get("name", ""),
                "bio":          profile.get("bio", ""),
                "location":     profile.get("location", ""),
                "email":        profile.get("email", ""),
                "website":      profile.get("blog", ""),
                "followers":    profile.get("followers", 0),
                "following":    profile.get("following", 0),
                "public_repos": profile.get("public_repos", 0),
                "github_url":   f"https://github.com/{username}",
            },
            "activity": activity,
            "skills_summary": {
                "top_languages":    top_languages,
                "language_counts":  dict(lang_counter.most_common(15)),
                "skills_from_code": list(set(all_skills)),
            },
            "repository_summary": {
                "total_repositories": len(repo_details),
                "non_fork_repos":     len(active_repos),
                "total_stars":        sum(r["stars"] for r in repo_details),
                "top_repositories":   [
                    {"name": r["name"], "stars": r["stars"],
                     "languages": r["languages"], "description": r["description"]}
                    for r in top_repos
                ],
            },
            "repositories":    repo_details,
            "final_assessment": final_assessment,
        }

    # ── Full pipeline ─────────────────────────────────────────────────────────

    def run(self):
        """Parse resume + scrape GitHub. Returns combined JSON."""
        resume_data = self.parse_resume()
        github_data = self.analyze_github_profile()
        return {"resume": resume_data, "github": github_data}


# ── CLI / main ────────────────────────────────────────────────────────────────

def main():
    agent         = Agent1("a.pdf")
    combined_text = agent.display()
    github_link   = agent.extract_github_link(combined_text)
    github_result = agent.analyze_github_profile() if github_link else {}

    final = {
        "resume_text": combined_text,
        "github":      github_result,
    }

    print(json.dumps(final, indent=2))
    return final

if __name__ == "__main__":
    main()