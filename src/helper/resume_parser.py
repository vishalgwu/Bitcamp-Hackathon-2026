import pdfplumber
import docx
import json
import os
import io
import re
import time
import fitz  # PyMuPDF
from google import genai
from dotenv import load_dotenv

load_dotenv()

client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))

# ─── Extract Hyperlinks from PDF ─────────────────────────

def extract_hyperlinks_from_pdf(file_bytes):
    """
    Uses PyMuPDF to extract all clickable hyperlinks.
    Correctly handles mailto: links as emails not portfolio URLs.
    """
    links = {
        "github_url":    None,
        "linkedin_url":  None,
        "portfolio_url": None,
        "email":         None,
        "all_links":     []
    }

    try:
        pdf = fitz.open(stream=file_bytes, filetype="pdf")

        for page in pdf:
            for link in page.get_links():
                url = link.get("uri", "")
                if not url:
                    continue

                links["all_links"].append(url)

                # mailto → extract email not portfolio
                if url.startswith("mailto:"):
                    email = url.replace("mailto:", "").strip()
                    if not links["email"]:
                        links["email"] = email
                        print(f"  📧 Email found in hyperlink: {email}")
                    continue

                # GitHub
                if "github.com" in url and not links["github_url"]:
                    parts = url.rstrip("/").split("/")
                    links["github_url"] = "/".join(parts[:4])
                    print(f"  🐙 GitHub URL found: {links['github_url']}")
                    continue

                # LinkedIn
                if "linkedin.com/in/" in url and not links["linkedin_url"]:
                    links["linkedin_url"] = url
                    print(f"  💼 LinkedIn URL found: {url}")
                    continue

                # Portfolio — anything else starting with http
                if (
                    url.startswith("http") and
                    "github.com" not in url and
                    "linkedin.com" not in url and
                    not links["portfolio_url"]
                ):
                    links["portfolio_url"] = url
                    print(f"  🌐 Portfolio URL found: {url}")

        pdf.close()
        print(f"  🔗 All hyperlinks: {links['all_links']}")

    except Exception as e:
        print(f"  ⚠️  Could not extract hyperlinks from PDF: {e}")

    return links

# ─── Extract Hyperlinks from DOCX ────────────────────────

def extract_hyperlinks_from_docx(file_bytes):
    """
    Extracts hyperlinks from DOCX relationships.
    Correctly handles mailto: links as emails.
    """
    links = {
        "github_url":    None,
        "linkedin_url":  None,
        "portfolio_url": None,
        "email":         None,
        "all_links":     []
    }

    try:
        doc = docx.Document(io.BytesIO(file_bytes))

        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype:
                url = rel._target
                if not url:
                    continue

                links["all_links"].append(url)

                # mailto → email
                if url.startswith("mailto:"):
                    email = url.replace("mailto:", "").strip()
                    if not links["email"]:
                        links["email"] = email
                        print(f"  📧 Email found in hyperlink: {email}")
                    continue

                # GitHub
                if "github.com" in url and not links["github_url"]:
                    parts = url.rstrip("/").split("/")
                    links["github_url"] = "/".join(parts[:4])
                    print(f"  🐙 GitHub URL found: {links['github_url']}")
                    continue

                # LinkedIn
                if "linkedin.com/in/" in url and not links["linkedin_url"]:
                    links["linkedin_url"] = url
                    print(f"  💼 LinkedIn URL found: {url}")
                    continue

                # Portfolio
                if (
                    url.startswith("http") and
                    "github.com" not in url and
                    "linkedin.com" not in url and
                    not links["portfolio_url"]
                ):
                    links["portfolio_url"] = url
                    print(f"  🌐 Portfolio URL found: {url}")

        print(f"  🔗 All hyperlinks: {links['all_links']}")

    except Exception as e:
        print(f"  ⚠️  Could not extract hyperlinks from DOCX: {e}")

    return links

# ─── Text Extraction ──────────────────────────────────────

def extract_text_from_pdf_bytes(file_bytes):
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_docx_bytes(file_bytes):
    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_file(file_bytes, filename):
    if filename.endswith(".pdf"):
        return extract_text_from_pdf_bytes(file_bytes)
    elif filename.endswith(".docx"):
        return extract_text_from_docx_bytes(file_bytes)
    else:
        raise ValueError("Unsupported format. Please upload PDF or DOCX.")

# ─── Gemini Parsing ───────────────────────────────────────

def parse_with_gemini(text):
    prompt = f"""
    Extract the following information from this resume text.
    Return ONLY a valid JSON object — no markdown, no extra text.
    
    {{
        "name": "",
        "email": "",
        "phone": "",
        "location": "",
        "summary": "Write a 2-3 sentence professional summary based on the resume",
        "github_url": "extract github url if present as plain text else empty string",
        "linkedin_url": "extract linkedin url if present as plain text else empty string",
        "portfolio_url": "extract any personal website or portfolio url if present as plain text else empty string",
        "skills": [],
        "experience": [
            {{
                "company": "",
                "role": "",
                "duration": "",
                "location": "",
                "description": ""
            }}
        ],
        "education": [
            {{
                "institution": "",
                "degree": "",
                "field": "",
                "year": ""
            }}
        ],
        "projects": [
            {{
                "name": "",
                "description": "",
                "technologies": []
            }}
        ],
        "certifications": []
    }}
    
    Resume text:
    {text}
    """

    models_to_try = [
        "models/gemini-2.0-flash",
        "models/gemini-2.5-flash",
        "models/gemini-2.0-flash-lite",
        "models/gemini-2.5-flash-lite",
    ]

    for model_name in models_to_try:
        for attempt in range(3):
            try:
                print(f"  Trying {model_name} (attempt {attempt + 1}/3)...")

                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt
                )

                # Token usage
                usage         = response.usage_metadata
                input_tokens  = usage.prompt_token_count
                output_tokens = usage.candidates_token_count
                total_tokens  = usage.total_token_count

                print(f"  ┌─────────────────────────────────────┐")
                print(f"  │         TOKEN USAGE SUMMARY          │")
                print(f"  ├─────────────────────────────────────┤")
                print(f"  │ Model         : {model_name:<21}│")
                print(f"  │ Input tokens  : {input_tokens:<21}│")
                print(f"  │ Output tokens : {output_tokens:<21}│")
                print(f"  │ Total tokens  : {total_tokens:<21}│")
                print(f"  └─────────────────────────────────────┘")

                raw = response.text.strip()
                raw = raw.replace("```json", "").replace("```", "").strip()
                result = json.loads(raw)
                print(f"  ✅ Success with {model_name}")
                return result

            except Exception as e:
                error_str = str(e)

                if "429" in error_str:
                    print(f"  ❌ Quota exhausted on {model_name} — trying next model...")
                    break

                elif "503" in error_str or "UNAVAILABLE" in error_str:
                    if attempt < 2:
                        wait = 30 * (attempt + 1)
                        print(f"  ⚠️  {model_name} unavailable. Waiting {wait}s before retry...")
                        time.sleep(wait)
                    else:
                        print(f"  ❌ {model_name} still unavailable — trying next model...")
                        break

                elif "404" in error_str or "NOT_FOUND" in error_str:
                    print(f"  ❌ {model_name} not found — trying next model...")
                    break

                else:
                    raise e

    raise Exception(
        "All models failed. Please wait a few minutes and try again."
    )

# ─── Merge Hyperlinks into Resume Data ───────────────────

def merge_links(resume_data, hyperlinks):
    """
    Hyperlinks from PDF/DOCX take priority over Gemini.
    mailto: links are treated as email, never as portfolio.
    """

    # Email from mailto takes priority
    if hyperlinks.get("email"):
        resume_data["email"] = hyperlinks["email"]
        print(f"  ✅ Email from mailto hyperlink: {hyperlinks['email']}")
    elif resume_data.get("email"):
        print(f"  ✅ Email from Gemini: {resume_data['email']}")
    else:
        print(f"  ⚠️  No email found")

    # GitHub
    if hyperlinks.get("github_url"):
        resume_data["github_url"] = hyperlinks["github_url"]
        print(f"  ✅ GitHub URL from hyperlink: {hyperlinks['github_url']}")
    elif resume_data.get("github_url"):
        print(f"  ✅ GitHub URL from Gemini: {resume_data['github_url']}")
    else:
        print(f"  ⚠️  No GitHub URL found")

    # LinkedIn
    if hyperlinks.get("linkedin_url"):
        resume_data["linkedin_url"] = hyperlinks["linkedin_url"]
        print(f"  ✅ LinkedIn URL from hyperlink: {hyperlinks['linkedin_url']}")
    elif resume_data.get("linkedin_url"):
        print(f"  ✅ LinkedIn URL from Gemini: {resume_data['linkedin_url']}")
    else:
        print(f"  ⚠️  No LinkedIn URL found")

    # Portfolio — never set if it's an email or mailto
    if hyperlinks.get("portfolio_url"):
        resume_data["portfolio_url"] = hyperlinks["portfolio_url"]
        print(f"  ✅ Portfolio URL from hyperlink: {hyperlinks['portfolio_url']}")
    elif resume_data.get("portfolio_url"):
        port = resume_data["portfolio_url"]
        # Fix if Gemini extracted email as portfolio
        if port and port.startswith("mailto:"):
            email = port.replace("mailto:", "").strip()
            print(f"  ⚠️  Gemini set mailto as portfolio — moving to email")
            if not resume_data.get("email"):
                resume_data["email"] = email
            resume_data["portfolio_url"] = None
        elif port and re.match(r'^[^@]+@[^@]+\.[^@]+$', port):
            print(f"  ⚠️  Gemini set email address as portfolio — moving to email")
            if not resume_data.get("email"):
                resume_data["email"] = port
            resume_data["portfolio_url"] = None
        else:
            print(f"  ✅ Portfolio URL from Gemini: {port}")
    else:
        print(f"  ⚠️  No portfolio URL found")

    # Store all links found
    resume_data["all_links"] = hyperlinks.get("all_links", [])

    return resume_data

# ─── Main Parse Functions ─────────────────────────────────

def parse_resume_from_upload(file_bytes, filename):
    """
    Parse from uploaded PDF or DOCX.
    Extracts text via pdfplumber/docx
    and hyperlinks via PyMuPDF/docx relationships.
    """
    print(f"  Parsing uploaded file: {filename}")

    # Extract text for Gemini
    text = extract_text_from_file(file_bytes, filename)

    # Extract hyperlinks
    print("  Extracting hyperlinks from file...")
    if filename.endswith(".pdf"):
        hyperlinks = extract_hyperlinks_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        hyperlinks = extract_hyperlinks_from_docx(file_bytes)
    else:
        hyperlinks = {}

    # Parse text with Gemini
    resume_data = parse_with_gemini(text)

    # Merge — hyperlinks take priority
    print("  Merging hyperlinks with Gemini output...")
    resume_data = merge_links(resume_data, hyperlinks)

    print("  ✅ Resume parsed successfully")
    return resume_data

def parse_resume_from_text(text):
    """
    Parse from pasted text.
    No hyperlinks possible — Gemini only.
    """
    print("  Parsing pasted resume text...")
    resume_data = parse_with_gemini(text)

    # Still check if Gemini set email as portfolio
    if resume_data.get("portfolio_url"):
        port = resume_data["portfolio_url"]
        if port and port.startswith("mailto:"):
            email = port.replace("mailto:", "").strip()
            if not resume_data.get("email"):
                resume_data["email"] = email
            resume_data["portfolio_url"] = None
            print(f"  ⚠️  Fixed: moved mailto to email field")
        elif port and re.match(r'^[^@]+@[^@]+\.[^@]+$', port):
            if not resume_data.get("email"):
                resume_data["email"] = port
            resume_data["portfolio_url"] = None
            print(f"  ⚠️  Fixed: moved email address to email field")

    print("  ✅ Resume parsed successfully")
    return resume_data

def parse_resume(source, filename=None):
    """
    Master function:
    - source = bytes → uploaded file (PDF/DOCX)
    - source = str   → pasted text
    """
    if isinstance(source, bytes):
        if not filename:
            raise ValueError("filename required for uploaded files")
        return parse_resume_from_upload(source, filename)
    elif isinstance(source, str):
        return parse_resume_from_text(source)
    else:
        raise ValueError("source must be file bytes or text string")
        