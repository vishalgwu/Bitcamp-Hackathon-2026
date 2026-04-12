"""
agent1.py — Resume Parser + GitHub Scraper
"""

import os, io, json, time, base64, threading, concurrent.futures, requests
from google import genai
from PyPDF2 import PdfReader
from dotenv import load_dotenv

load_dotenv()

# ── API Keys ──────────────────────────────────────────────────────────────────
GEMINI_API_KEY  = os.environ.get("GEMINI_API_KEY")
GEMINI_CODE_KEY = os.environ.get("GEMINI_CODE_KEY") or os.environ.get("GEMINI_API_KEY")
GITHUB_TOKEN    = os.environ.get("GITHUB_TOKEN")


class Agent1:

    MODELS = [
        "gemini-2.5-flash",
        "gemini-2.5-pro",
        "gemini-2.0-flash",
        "gemini-2.0-flash-lite",
    ]

    CODE_EXTENSIONS = {
        ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".cpp", ".c",
        ".cs", ".go", ".rs", ".swift", ".kt", ".sql", ".ipynb", ".rb"
    }

    SKIP_PATTERNS = [
        "package-lock.json", "yarn.lock", ".gitignore", "node_modules",
        "dist/", "build/", "__pycache__", ".min.js", "vendor/"
    ]

    def __init__(self):
        # Resume parsing + portfolio scraping
        self.gemini      = genai.Client(api_key=GEMINI_API_KEY)
        # GitHub code analysis — uses separate quota key
        self.code_gemini = genai.Client(api_key=GEMINI_CODE_KEY)
        self.gh_headers  = {
            "Authorization": f"token {GITHUB_TOKEN}",
            "Accept":        "application/vnd.github.v3+json"
        }

    # =========================================================================
    # SECTION 1 — RESUME PARSER
    # =========================================================================

    def _extract_text(self, file_bytes, filename):
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages)
            except:
                try:
                    reader = PdfReader(io.BytesIO(file_bytes))
                    return "\n".join(p.extract_text() or "" for p in reader.pages)
                except:
                    return ""
        elif ext in [".docx", ".doc"]:
            try:
                from docx import Document
                doc = Document(io.BytesIO(file_bytes))
                return "\n".join(p.text for p in doc.paragraphs)
            except:
                return ""
        return file_bytes.decode("utf-8", errors="replace")

    def _extract_hyperlinks(self, file_bytes, filename):
        ext   = os.path.splitext(filename)[1].lower()
        links = []

        if ext == ".pdf":
            # ── Method 1: PyMuPDF (fitz) — most reliable for PDFs ────────────
            try:
                import fitz
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page in doc:
                    for link in page.get_links():
                        uri = link.get("uri", "")
                        if uri and isinstance(uri, str) and uri.startswith("http"):
                            links.append(uri)
                doc.close()
            except Exception:
                pass

            # ── Method 2: pdfplumber annotations fallback ─────────────────────
            if not links:
                try:
                    import pdfplumber
                    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                        for page in pdf.pages:
                            for annot in (page.annots or []):
                                uri = annot.get("uri") or annot.get("A", {}).get("URI")
                                if uri and isinstance(uri, str) and uri.startswith("http"):
                                    links.append(uri)
                except Exception:
                    pass

        elif ext in [".docx", ".doc"]:
            try:
                from docx import Document
                doc = Document(io.BytesIO(file_bytes))
                for rel in doc.part.rels.values():
                    if "hyperlink" in rel.reltype:
                        url = rel._target
                        if isinstance(url, str) and url.startswith("http"):
                            links.append(url)
            except Exception:
                pass

        # Deduplicate + extract ONLY github profile URLs (not repo URLs)
        seen    = set()
        cleaned = []
        for url in links:
            url = url.strip().rstrip("/")
            if url not in seen:
                seen.add(url)
                cleaned.append(url)

        return cleaned

    def _extract_github_url(self, hyperlinks):
        """
        From a list of hyperlinks, find the GitHub PROFILE url.
        Profile:  github.com/username         (1 path segment) ✅
        Repo:     github.com/username/repo    (2 segments)     ❌ — extract username
        """
        for url in hyperlinks:
            if "github.com" not in url:
                continue
            # Strip protocol + domain
            path = url.lower()
            for prefix in ["https://github.com/", "http://github.com/",
                           "https://www.github.com/"]:
                if path.startswith(prefix):
                    path = url[len(prefix):]
                    break
            parts = [p for p in path.rstrip("/").split("/") if p]
            if len(parts) == 1:
                # Already a profile URL
                return f"https://github.com/{parts[0]}"
            elif len(parts) >= 2:
                # Repo URL — extract just the username part
                return f"https://github.com/{parts[0]}"
        return ""

    def _parse_with_gemini(self, text, hyperlinks):
        links_str = "\n".join(hyperlinks) if hyperlinks else "None found"
        prompt = f"""
Extract all structured information from this resume.
Pay special attention to the hyperlinks list — these are ACTUAL URLs embedded in the document.
Return ONLY valid JSON — no markdown, no extra text:

{{
    "name": "", "email": "", "phone": "", "location": "", "summary": "",
    "github_url": "", "linkedin_url": "", "portfolio_url": "",
    "skills": [],
    "experience": [{{"role": "", "company": "", "duration": "", "location": "", "description": ""}}],
    "education": [{{"degree": "", "field": "", "institution": "", "year": ""}}],
    "certifications": [],
    "projects": [{{"name": "", "description": "", "technologies": []}}]
}}

Hyperlinks found in document:
{links_str}

Resume text:
{text[:6000]}
"""
        result = self._gemini_json(prompt)

        # Override github_url with the one extracted directly from hyperlinks
        # — more reliable than Gemini's extraction from text
        github_from_links = self._extract_github_url(hyperlinks)
        if github_from_links:
            result["github_url"] = github_from_links

        return result

    def parse_resume(self, file_bytes, filename):
        print(f"\n📄 Parsing resume: {filename}")
        text  = self._extract_text(file_bytes, filename)
        print(f"  ✅ Extracted {len(text)} chars")
        links = self._extract_hyperlinks(file_bytes, filename)
        print(f"  ✅ Found {len(links)} hyperlinks: {links}")
        result = self._parse_with_gemini(text, links)
        print(f"  ✅ Parsed successfully")
        return result

    # =========================================================================
    # SECTION 2 — GITHUB SCRAPER
    # =========================================================================

    def _gh_get(self, url, as_list=False):
        try:
            r = requests.get(url, headers=self.gh_headers, timeout=10)
            data = r.json()
            if r.status_code == 200:
                return data
        except Exception:
            pass
        return [] if as_list else {}

    def _get_languages(self, username, repo_name):
        data = self._gh_get(f"https://api.github.com/repos/{username}/{repo_name}/languages")
        return list(data.keys()) if isinstance(data, dict) else []

    def _get_readme(self, username, repo_name):
        data = self._gh_get(f"https://api.github.com/repos/{username}/{repo_name}/readme")
        if isinstance(data, dict) and data.get("content"):
            try:
                return base64.b64decode(
                    data["content"].replace("\n", "")
                ).decode("utf-8", errors="replace")[:500]
            except:
                pass
        return ""

    def _get_commit_count(self, username, repo_name):
        try:
            r = requests.get(
                f"https://api.github.com/repos/{username}/{repo_name}/commits?per_page=1",
                headers=self.gh_headers, timeout=10
            )
            if r.status_code == 200:
                if "Link" in r.headers:
                    try:
                        last = r.headers["Link"].split(",")[-1]
                        return int(last.split("page=")[-1].split(">")[0])
                    except:
                        pass
                return len(r.json())
        except:
            pass
        return 0

    def _fetch_repo_code(self, username, repo_name):
        # Try multiple branches — repos may not use main/master
        tree = []
        for branch in ["main", "master", "HEAD", "dev", "develop"]:
            try:
                r = requests.get(
                    f"https://api.github.com/repos/{username}/{repo_name}/git/trees/{branch}?recursive=1",
                    headers=self.gh_headers, timeout=15
                )
                if r.status_code == 200:
                    data = r.json()
                    tree = data.get("tree", [])
                    if tree:
                        break
            except:
                pass

        files = []
        for item in tree:
            if item.get("type") != "blob":
                continue
            path = item.get("path", "")
            if any(s in path for s in self.SKIP_PATTERNS):
                continue
            if os.path.splitext(path)[1].lower() not in self.CODE_EXTENSIONS:
                continue
            if item.get("size", 0) > 100000:
                continue
            files.append(path)
            if len(files) >= 5:
                break

        code_samples = {}
        for path in files:
            content = ""

            # Method 1 — GitHub Contents API (base64)
            try:
                r = requests.get(
                    f"https://api.github.com/repos/{username}/{repo_name}/contents/{path}",
                    headers=self.gh_headers, timeout=10
                )
                if r.status_code == 200:
                    data = r.json()
                    if isinstance(data, dict):
                        if data.get("encoding") == "base64" and data.get("content"):
                            content = base64.b64decode(
                                data["content"].replace("\n", "")
                            ).decode("utf-8", errors="replace")
                        elif data.get("download_url"):
                            # Some files return download_url instead of base64
                            r2 = requests.get(data["download_url"], timeout=10)
                            if r2.status_code == 200:
                                content = r2.text
            except:
                pass

            # Method 2 — raw.githubusercontent.com fallback
            if not content or not content.strip():
                for branch in ["main", "master", "HEAD"]:
                    try:
                        r = requests.get(
                            f"https://raw.githubusercontent.com/{username}/{repo_name}/{branch}/{path}",
                            headers={"Authorization": f"token {GITHUB_TOKEN}"},
                            timeout=10
                        )
                        if r.status_code == 200 and r.text.strip():
                            content = r.text
                            break
                    except:
                        pass

            if content and content.strip():
                code_samples[path] = content

        return repo_name, code_samples

    def _analyze_code(self, repo_name, code_samples):
        if not code_samples:
            return {}
        code_context = ""
        for path, content in code_samples.items():
            code_context += f"\n\n--- {path} ---\n{content[:1500]}"
        prompt = f"""
Analyze code from GitHub repo "{repo_name}".
Return ONLY valid JSON — no markdown:

{{
    "code_quality": {{"score": 0, "rating": "", "summary": ""}},
    "technical_complexity": {{"score": 0, "rating": "", "summary": ""}},
    "architecture_patterns": [],
    "skills_demonstrated": [],
    "best_practices_used": [],
    "areas_for_improvement": [],
    "overall_assessment": "",
    "notable_observations": []
}}

score: 1-10, rating: Beginner/Intermediate/Advanced/Expert

Code:
{code_context}
"""
        return self._code_gemini_json(prompt)

    def _assess_developer(self, repo_details, all_skills, all_languages):
        summary = f"""
Repos: {len(repo_details)}
Languages: {', '.join(all_languages[:10])}
Skills from code: {', '.join(list(set(all_skills))[:15])}
Top repos: {[r['name'] for r in sorted(repo_details, key=lambda x: x['stars'], reverse=True)[:5]]}
"""
        prompt = f"""
Assess the developer based on their GitHub profile summary.
Return ONLY valid JSON — no markdown:

{{
    "developer_level": "",
    "strengths": [],
    "weaknesses": [],
    "confidence_score": 0.0
}}

developer_level: Beginner / Intermediate / Advanced / Expert

Profile:
{summary}
"""
        return self._gemini_json(prompt)

    def scrape_github(self, github_url):
        if not github_url:
            print("  ⏭️  No GitHub URL — skipping")
            return {}

        username = github_url.rstrip("/").split("/")[-1]
        print(f"\n🐙 Scraping GitHub: {username}")

        profile = self._gh_get(f"https://api.github.com/users/{username}")
        if not profile or "message" in profile:
            print(f"  ❌ GitHub user not found: {username}")
            return {}
        
        # Use the correct casing from the API response
        username = profile.get("login", username)
        print(f"  ✅ Corrected username: {username}")

        public_repos = profile.get("public_repos", 0)
        if public_repos > 200:
            print(f"  ⚠️  Org account ({public_repos} repos) — skipping")
            return {}

        print(f"  ✅ Profile found — {public_repos} repos")

        repos_raw = self._gh_get(
            f"https://api.github.com/users/{username}/repos?per_page=100&sort=updated",
            as_list=True
        )
        print(f"  📦 Fetching metadata for {len(repos_raw)} repos...")

        repo_details = []
        for repo in repos_raw:
            name = repo["name"]
            repo_details.append({
                "name":           name,
                "description":    repo.get("description") or "",
                "languages":      self._get_languages(username, name),
                "stars":          repo.get("stargazers_count", 0),
                "forks":          repo.get("forks_count", 0),
                "commit_count":   self._get_commit_count(username, name),
                "last_updated":   repo.get("updated_at", ""),
                "topics":         repo.get("topics", []),
                "has_live_demo":  bool(repo.get("homepage")),
                "homepage":       repo.get("homepage", ""),
                "is_fork":        repo.get("fork", False),
                "readme_preview": self._get_readme(username, name),
                "code_samples":   {},
                "code_insights":  {}
            })

        # Code analysis — sequential, all non-fork repos
        non_forks = sorted(
            [r for r in repo_details if not r["is_fork"]],
            key=lambda x: x["commit_count"], reverse=True
        )
        print(f"  🔍 Fetching + analyzing all {len(non_forks)} repos sequentially...")

        repos_analyzed = 0
        for i, r in enumerate(non_forks):
            repo_name = r["name"]
            print(f"  [{i+1}/{len(non_forks)}] {repo_name}")

            # Fetch code
            _, code_samples = self._fetch_repo_code(username, repo_name)
            if not code_samples:
                print(f"    ⚠️  No code files found")
                continue

            # Analyze with Gemini code key
            insights = self._analyze_code(repo_name, code_samples)

            # Store results
            for repo in repo_details:
                if repo["name"] == repo_name:
                    repo["code_samples"]  = {k: v[:500] for k, v in code_samples.items()}
                    repo["code_insights"] = insights
                    if insights:
                        repos_analyzed += 1

            # Delay between repos — shorter now since quota is handled inside _call_gemini
            if i < len(non_forks) - 1:
                delay = 5 if (i + 1) % 3 == 0 else 2
                time.sleep(delay)

        # Aggregate
        all_languages, all_skills, all_patterns, all_practices = [], [], [], []
        for repo in repo_details:
            all_languages += repo["languages"]
            ins = repo.get("code_insights", {})
            all_skills    += ins.get("skills_demonstrated", [])
            all_patterns  += ins.get("architecture_patterns", [])
            all_practices += ins.get("best_practices_used", [])

        all_languages = list(set(all_languages))
        top_repos     = sorted(repo_details, key=lambda x: x["stars"], reverse=True)[:5]
        final_assessment = self._assess_developer(repo_details, all_skills, all_languages)

        print(f"  ✅ GitHub done — {len(repo_details)} repos, {repos_analyzed} analyzed")

        return {
            "candidate_profile": {
                "github": {
                    "username":     username,
                    "name":         profile.get("name", ""),
                    "bio":          profile.get("bio", ""),
                    "location":     profile.get("location", ""),
                    "email":        profile.get("email", ""),
                    "website":      profile.get("blog", ""),
                    "followers":    profile.get("followers", 0),
                    "following":    profile.get("following", 0),
                    "public_repos": public_repos,
                    "github_url":   f"https://github.com/{username}",
                }
            },
            "skills_summary": {
                "all_languages":         all_languages,
                "skills_from_code":      list(set(all_skills)),
                "architecture_patterns": list(set(all_patterns)),
                "best_practices":        list(set(all_practices)),
            },
            "repository_summary": {
                "total_repositories":    len(repo_details),
                "analyzed_repositories": repos_analyzed,
                "top_repositories": [
                    {"name": r["name"], "stars": r["stars"],
                     "languages": r["languages"], "description": r["description"]}
                    for r in top_repos
                ],
            },
            "repositories":     repo_details,
            "final_assessment": final_assessment,
        }

    # =========================================================================
    # SECTION 3 — GEMINI HELPER
    # =========================================================================

    def _gemini_json(self, prompt):
        """Uses GEMINI_API_KEY — for resume parsing + portfolio + developer assessment"""
        return self._call_gemini(self.gemini, prompt, fallback_client=self.code_gemini)

    def _code_gemini_json(self, prompt):
        """Uses GEMINI_CODE_KEY — for GitHub code analysis only"""
        return self._call_gemini(self.code_gemini, prompt, fallback_client=self.gemini)

    def _call_gemini(self, client, prompt, fallback_client=None):
        """
        Shared Gemini caller.
        On quota hit → waits briefly → tries next model on same client.
        If ALL models on primary client exhausted → switches to fallback_client.
        This means GEMINI_CODE_KEY and GEMINI_API_KEY back each other up.
        """
        def repair_json(text):
            text = text.strip()
            last_valid = max(text.rfind(','), text.rfind('}'), text.rfind(']'))
            if last_valid > 0:
                text = text[:last_valid]
            open_braces   = text.count('{') - text.count('}')
            open_brackets = text.count('[') - text.count(']')
            text += ']' * open_brackets
            text += '}' * open_braces
            return text

        def try_client(c, label):
            for model in self.MODELS:
                try:
                    response = c.models.generate_content(
                        model=model, contents=prompt
                    )
                    raw = response.text.strip()
                    raw = raw.replace("```json", "").replace("```", "").strip()
                    try:
                        return json.loads(raw)
                    except json.JSONDecodeError:
                        print(f"  [{label}] JSON truncated from {model}, repairing...")
                        try:
                            return json.loads(repair_json(raw))
                        except:
                            print(f"  [{label}] Repair failed for {model}, trying next...")
                            continue
                except Exception as e:
                    err = str(e)
                    if "429" in err or "RESOURCE_EXHAUSTED" in err:
                        print(f"  [{label}] {model} quota hit — trying next model...")
                        time.sleep(5)   # short wait then immediately next model
                    elif "503" in err or "unavailable" in err.lower():
                        print(f"  [{label}] {model} unavailable — trying next model...")
                        time.sleep(3)
                    else:
                        print(f"  [{label}] {model} error: {err[:80]}")
                    continue
            return None

        # Try primary client first
        result = try_client(client, "primary")
        if result is not None:
            return result

        # All models on primary exhausted — switch to fallback key
        if fallback_client and fallback_client is not client:
            print(f"  Primary key exhausted — switching to fallback key...")
            result = try_client(fallback_client, "fallback")
            if result is not None:
                return result

        # Both keys exhausted — wait longer then try primary once more
        print(f"  Both keys exhausted — waiting 90s then retrying...")
        time.sleep(90)
        result = try_client(client, "retry")
        if result is not None:
            return result

        return {}

    # =========================================================================
    # SECTION 4 — MAIN ENTRY POINT
    # =========================================================================

    def run(self, file_bytes, filename):
        resume_data = self.parse_resume(file_bytes, filename)
        github_url  = resume_data.get("github_url", "")
        github_data = self.scrape_github(github_url) if github_url else {}
        return {"resume": resume_data, "github": github_data}


# ── CLI ───────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python agent1.py resume.pdf")
        sys.exit(1)
    with open(sys.argv[1], "rb") as f:
        data = f.read()
    result = Agent1().run(data, os.path.basename(sys.argv[1]))
    print(json.dumps(result, indent=2))